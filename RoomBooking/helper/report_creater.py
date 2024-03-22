from datetime import datetime

from docx import Document

from RoomBooking.models import Booking
from Booking.settings import BASE_DIR


def create_report(room_number=None, date_time_start=None, date_time_end=None) -> None:
    """
    Создать отчет

    :params room_number - номер комнаты
    :params date_time_start - начало временого диапозона бронирований
    :params date_time_end - конец временого диапозона бронирований
    """
    schedules = Booking.get_all_schedule()
    document = Document()
    document.add_heading('Расписание бронирования комнат', level=1).bold = True
    document.add_paragraph()

    if date_time_start:
        date_time_start = datetime.strptime(date_time_start, "%Y-%m-%dT%H:%M:%S")

    if date_time_end:
        date_time_end = datetime.strptime(date_time_end, "%Y-%m-%dT%H:%M:%S")

    for schedule in schedules:
        room_number_pass = room_number is None or schedule.room.number == room_number
        date_time_start_pass = (date_time_start is None or
                                date_time_start <= schedule.date_time_start < date_time_end)
        date_time_end_pass = (date_time_end is None or
                              date_time_start <= schedule.date_time_end < date_time_end)

        if room_number_pass and (date_time_start_pass or date_time_end_pass):
            document.add_heading("Номер комнаты: " + schedule.room.number, level=2).bold = True
            document.add_paragraph("Забронировал: " + schedule.user.username)
            document.add_paragraph("Дата(от): " + str(schedule.date_time_start))
            document.add_paragraph("Дата(до): " + str(schedule.date_time_end))
            document.add_paragraph("Цель бронирования: " + schedule.purpose)
            document.add_paragraph()

    report_save_path = f"{BASE_DIR}/RoomBooking/storage/booking_report/report.docx"
    document.save(report_save_path)
